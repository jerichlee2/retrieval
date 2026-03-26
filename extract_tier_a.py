#!/usr/bin/env python3
"""
Extract Tier A text-bearing files from a download manifest into a normalized JSONL corpus.

Tier A extensions:
    .txt, .pdf, .csv, .docx, .md, .html

Outputs:
    1. extracted_documents.jsonl  -> one JSON object per extracted source file
    2. extraction_report.csv      -> per-file success/failure report

Example:
    python extract_tier_a.py \
        --manifest download_manifest.csv \
        --output-jsonl extracted_documents_tier_a.jsonl \
        --report-csv extraction_report_tier_a.csv
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


TIER_A_EXTENSIONS = {".txt", ".pdf", ".csv", ".docx", ".md", ".html"}
SUCCESS_STATUSES = {"download", "export"}
TEXTUAL_EXTENSIONS = {".txt", ".md"}
CSV_EXTENSIONS = {".csv"}
HTML_EXTENSIONS = {".html"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract Tier A files into a JSONL corpus.")
    parser.add_argument(
        "--manifest",
        required=True,
        help="Path to download_manifest.csv",
    )
    parser.add_argument(
        "--output-jsonl",
        default="extracted_documents_tier_a.jsonl",
        help="Output JSONL path",
    )
    parser.add_argument(
        "--report-csv",
        default="extraction_report_tier_a.csv",
        help="Output extraction report CSV path",
    )
    parser.add_argument(
        "--base-dir",
        default=".",
        help=(
            "Base directory used to resolve relative local_path values from the manifest. "
            "Usually the directory where rag_downloads/ lives."
        ),
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Optional limit for quick testing.",
    )
    parser.add_argument(
        "--include-empty",
        action="store_true",
        help="Keep records even if extracted text is empty.",
    )
    return parser.parse_args()


def normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    return text.strip()


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def resolve_local_path(base_dir: Path, local_path_value: str) -> Path:
    p = Path(str(local_path_value))
    if p.is_absolute():
        return p
    return (base_dir / p).resolve()


def read_text_file(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "latin-1"]
    last_error: Optional[Exception] = None
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except Exception as exc:  # pragma: no cover - best effort fallback
            last_error = exc
    raise RuntimeError(f"Could not decode text file: {path}") from last_error


def extract_from_pdf(path: Path) -> Tuple[str, Dict[str, Any]]:
    reader = PdfReader(str(path))
    page_texts: List[str] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover
            page_text = f"[PDF extraction error on page {i}: {exc}]"
        page_text = normalize_whitespace(page_text)
        if page_text:
            page_texts.append(f"[Page {i}]\n{page_text}")
    return "\n\n".join(page_texts), {"page_count": len(reader.pages)}


def extract_from_docx(path: Path) -> Tuple[str, Dict[str, Any]]:
    doc = Document(str(path))
    parts: List[str] = []

    paragraphs = [normalize_whitespace(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]
    if paragraphs:
        parts.append("\n".join(paragraphs))

    table_count = 0
    for table in doc.tables:
        table_count += 1
        rows: List[str] = []
        for row in table.rows:
            cells = [normalize_whitespace(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table {table_count}]\n" + "\n".join(rows))

    return "\n\n".join(parts), {"table_count": table_count}


def extract_from_csv(path: Path) -> Tuple[str, Dict[str, Any]]:
    raw = read_text_file(path)
    sample = raw[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except Exception:
        dialect = csv.excel

    rows_out: List[str] = []
    row_count = 0
    reader = csv.reader(raw.splitlines(), dialect)
    for row in reader:
        row_count += 1
        clean = [normalize_whitespace(cell) for cell in row]
        rows_out.append(" | ".join(clean))

    return "\n".join(rows_out), {"row_count": row_count}


def extract_from_html(path: Path) -> Tuple[str, Dict[str, Any]]:
    raw = read_text_file(path)
    soup = BeautifulSoup(raw, "html.parser")
    title = soup.title.get_text(strip=True) if soup.title else ""
    text = soup.get_text("\n")
    text = html.unescape(text)
    text = normalize_whitespace(text)
    return text, {"html_title": title}


def extract_text(path: Path, extension: str) -> Tuple[str, Dict[str, Any]]:
    if extension in TEXTUAL_EXTENSIONS:
        return read_text_file(path), {}
    if extension in PDF_EXTENSIONS:
        return extract_from_pdf(path)
    if extension in CSV_EXTENSIONS:
        return extract_from_csv(path)
    if extension in DOCX_EXTENSIONS:
        return extract_from_docx(path)
    if extension in HTML_EXTENSIONS:
        return extract_from_html(path)
    raise ValueError(f"Unsupported Tier A extension: {extension}")


def build_record(row: pd.Series, local_file_path: Path, extension: str, text: str, extra_meta: Dict[str, Any]) -> Dict[str, Any]:
    text = normalize_whitespace(text)
    return {
        "source_id": row.get("id", ""),
        "source_name": row.get("name", ""),
        "source_path": row.get("path", ""),
        "local_path": str(local_file_path),
        "relative_local_path": row.get("local_path", ""),
        "file_type": extension.lstrip("."),
        "mime_type": row.get("mimeType", ""),
        "web_view_link": row.get("webViewLink", ""),
        "download_mode": row.get("mode", ""),
        "manifest_status": row.get("status", ""),
        "char_count": len(text),
        "word_count": count_words(text),
        "text_sha256": sha256_text(text),
        "text": text,
        "metadata": extra_meta,
    }


def iter_tier_a_rows(df: pd.DataFrame) -> Iterable[pd.Series]:
    for _, row in df.iterrows():
        status = str(row.get("status", "")).strip().lower()
        if status not in SUCCESS_STATUSES:
            continue

        local_path_value = str(row.get("local_path", "")).strip()
        extension = Path(local_path_value).suffix.lower()
        if extension in TIER_A_EXTENSIONS:
            yield row


def main() -> None:
    args = parse_args()
    base_dir = Path(args.base_dir).resolve()

    manifest_path = Path(args.manifest)
    if not manifest_path.exists():
        raise FileNotFoundError(f"Manifest not found: {manifest_path}")

    df = pd.read_csv(manifest_path)
    rows = list(iter_tier_a_rows(df))
    if args.limit is not None:
        rows = rows[: args.limit]

    output_jsonl = Path(args.output_jsonl)
    report_csv = Path(args.report_csv)
    output_jsonl.parent.mkdir(parents=True, exist_ok=True)
    report_csv.parent.mkdir(parents=True, exist_ok=True)

    report_rows: List[Dict[str, Any]] = []
    success_count = 0
    error_count = 0
    skipped_empty_count = 0

    with output_jsonl.open("w", encoding="utf-8") as fout:
        total = len(rows)
        for idx, row in enumerate(rows, start=1):
            rel_local_path = str(row.get("local_path", "")).strip()
            extension = Path(rel_local_path).suffix.lower()
            local_file_path = resolve_local_path(base_dir, rel_local_path)

            try:
                if not local_file_path.exists():
                    raise FileNotFoundError(f"Local file not found: {local_file_path}")

                text, extra_meta = extract_text(local_file_path, extension)
                text = normalize_whitespace(text)

                if not text and not args.include_empty:
                    skipped_empty_count += 1
                    report_rows.append({
                        "index": idx,
                        "source_id": row.get("id", ""),
                        "source_path": row.get("path", ""),
                        "local_path": str(local_file_path),
                        "extension": extension,
                        "status": "skipped_empty",
                        "char_count": 0,
                        "word_count": 0,
                        "error": "",
                    })
                    print(f"[{idx}/{total}] SKIP empty -> {rel_local_path}", flush=True)
                    continue

                record = build_record(row, local_file_path, extension, text, extra_meta)
                fout.write(json.dumps(record, ensure_ascii=False) + "\n")
                success_count += 1

                report_rows.append({
                    "index": idx,
                    "source_id": row.get("id", ""),
                    "source_path": row.get("path", ""),
                    "local_path": str(local_file_path),
                    "extension": extension,
                    "status": "ok",
                    "char_count": record["char_count"],
                    "word_count": record["word_count"],
                    "error": "",
                })
                print(
                    f"[{idx}/{total}] OK {extension:>5} -> {rel_local_path} "
                    f"({record['char_count']} chars)",
                    flush=True,
                )

            except Exception as exc:
                error_count += 1
                report_rows.append({
                    "index": idx,
                    "source_id": row.get("id", ""),
                    "source_path": row.get("path", ""),
                    "local_path": str(local_file_path),
                    "extension": extension,
                    "status": "error",
                    "char_count": 0,
                    "word_count": 0,
                    "error": str(exc),
                })
                print(f"[{idx}/{total}] ERROR {extension:>5} -> {rel_local_path} | {exc}", flush=True)

    report_df = pd.DataFrame(report_rows)
    report_df.to_csv(report_csv, index=False)

    print("\nDone.", flush=True)
    print(f"  Output JSONL:      {output_jsonl}", flush=True)
    print(f"  Extraction report: {report_csv}", flush=True)
    print(f"  Successes:         {success_count}", flush=True)
    print(f"  Empty skipped:     {skipped_empty_count}", flush=True)
    print(f"  Errors:            {error_count}", flush=True)


if __name__ == "__main__":
    main()
